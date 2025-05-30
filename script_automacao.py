import streamlit as st
import csv
import os
from docx import Document
from docx.shared import Pt
import zipfile
import re
import pythoncom 
from docx2pdf import convert

st.set_page_config(page_title="Automatização Contratos")
st.title("Gerador de Contratos PDF")

arquivo_docx = st.file_uploader("Arquivo .docx", type="docx")
arquivo_csv = st.file_uploader("Arquivo .csv", type="csv")
caminho_zip = st.text_input("Caminho completo para salvar o arquivo ZIP", placeholder=r"C:\Users\user\Desktop\contratos.zip")

if st.button("Gerar"):
    if arquivo_docx and arquivo_csv and caminho_zip:
        conteudo_csv = arquivo_csv.read().decode("utf-8")
        leitor = csv.reader(conteudo_csv.splitlines())
        next(leitor)
        pessoas = list(leitor)

        pdfs = []

        conteudo_docx = arquivo_docx.read()
        with open("modelo_temp.docx", "wb") as f:
            f.write(conteudo_docx)

        pythoncom.CoInitialize()

        for i, pessoa in enumerate(pessoas):
            nome = re.sub(r"[\[\]']", "", pessoa[0])
            doc = Document("modelo_temp.docx")
            paragrafo = doc.paragraphs[0]
            paragrafo.clear()
            texto = paragrafo.add_run(nome)
            texto.bold = True
            texto.font.size = Pt(20)

            nome_docx = f"contrato{i+1}.docx"
            nome_pdf = f"contrato{i+1}.pdf"
            doc.save(nome_docx)
            convert(nome_docx, nome_pdf)
            os.remove(nome_docx)
            pdfs.append(nome_pdf)

        with zipfile.ZipFile(caminho_zip, "w", zipfile.ZIP_DEFLATED) as zipf:
            for pdf in pdfs:
                zipf.write(pdf)
                os.remove(pdf)

        st.success(f"ZIP criado em: {caminho_zip}")
        os.remove("modelo_temp.docx")
    else:
        st.warning("Preencha tudo antes de gerar.")
